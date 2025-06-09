# rag-studio/backend/app/utils/file_parser.py
"""
파일 파서 유틸리티

다양한 형식의 문서를 파싱하여 색인 가능한 형태로 변환합니다.
"""

import os
import json
import csv
import io
from typing import List, Dict, Any, Optional
from pathlib import Path
from datetime import datetime
import hashlib
import mimetypes

import aiofiles
import pandas as pd
from pypdf import PdfReader
import docx
import chardet

from app.utils.logger import logger
from app.schemas.opensearch import DocumentInput


async def parse_document_file(
    file_path: Path,
    file_extension: str,
    source: str = "upload"
) -> List[DocumentInput]:
    """
    문서 파일을 파싱하여 DocumentInput 리스트로 변환
    
    Args:
        file_path: 파일 경로
        file_extension: 파일 확장자
        source: 문서 출처
        
    Returns:
        List[DocumentInput]: 파싱된 문서 리스트
    """
    try:
        file_extension = file_extension.lower()
        
        if file_extension == "txt":
            documents = await _parse_text_file(file_path, source)
        elif file_extension == "pdf":
            documents = await _parse_pdf_file(file_path, source)
        elif file_extension == "docx":
            documents = await _parse_docx_file(file_path, source)
        elif file_extension == "csv":
            documents = await _parse_csv_file(file_path, source)
        elif file_extension == "json":
            documents = await _parse_json_file(file_path, source)
        else:
            raise ValueError(f"Unsupported file extension: {file_extension}")
        
        logger.info(f"파일 파싱 완료: {file_path.name}, 문서 수: {len(documents)}")
        return documents
        
    except Exception as e:
        logger.error(f"파일 파싱 중 오류: {file_path.name}, {str(e)}")
        raise


async def _parse_text_file(file_path: Path, source: str) -> List[DocumentInput]:
    """
    텍스트 파일 파싱
    
    Args:
        file_path: 파일 경로
        source: 문서 출처
        
    Returns:
        List[DocumentInput]: 파싱된 문서
    """
    documents = []
    
    # 인코딩 감지
    async with aiofiles.open(file_path, 'rb') as f:
        raw_data = await f.read()
        detected = chardet.detect(raw_data)
        encoding = detected['encoding'] or 'utf-8'
    
    # 텍스트 읽기
    async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
        content = await f.read()
    
    # 문서 ID 생성
    doc_id = _generate_document_id(file_path.name, content)
    
    # DocumentInput 생성
    document = DocumentInput(
        document_id=doc_id,
        title=file_path.stem,  # 파일명에서 확장자 제거
        content=content,
        source=source,
        metadata={
            "file_name": file_path.name,
            "file_size": file_path.stat().st_size,
            "encoding": encoding,
            "created_at": datetime.utcnow().isoformat()
        }
    )
    
    documents.append(document)
    return documents


async def _parse_pdf_file(file_path: Path, source: str) -> List[DocumentInput]:
    """
    PDF 파일 파싱
    
    Args:
        file_path: 파일 경로
        source: 문서 출처
        
    Returns:
        List[DocumentInput]: 파싱된 문서
    """
    documents = []
    
    # PDF 읽기
    reader = PdfReader(str(file_path))
    
    # 전체 텍스트 추출
    full_text = ""
    page_texts = []
    
    for page_num, page in enumerate(reader.pages, 1):
        page_text = page.extract_text()
        page_texts.append(page_text)
        full_text += page_text + "\n\n"
    
    # 문서 ID 생성
    doc_id = _generate_document_id(file_path.name, full_text)
    
    # 메타데이터 수집
    metadata = {
        "file_name": file_path.name,
        "file_size": file_path.stat().st_size,
        "page_count": len(reader.pages),
        "created_at": datetime.utcnow().isoformat()
    }
    
    # PDF 메타데이터 추가
    if reader.metadata:
        if reader.metadata.title:
            metadata["pdf_title"] = reader.metadata.title
        if reader.metadata.author:
            metadata["pdf_author"] = reader.metadata.author
        if reader.metadata.subject:
            metadata["pdf_subject"] = reader.metadata.subject
        if reader.metadata.creator:
            metadata["pdf_creator"] = reader.metadata.creator
    
    # 전체 문서를 하나로 저장
    document = DocumentInput(
        document_id=doc_id,
        title=metadata.get("pdf_title", file_path.stem),
        content=full_text.strip(),
        source=source,
        metadata=metadata
    )
    
    documents.append(document)
    
    # 옵션: 각 페이지를 별도 문서로 저장
    # for page_num, page_text in enumerate(page_texts, 1):
    #     if page_text.strip():
    #         page_doc = DocumentInput(
    #             document_id=f"{doc_id}_page_{page_num}",
    #             title=f"{document.title} - Page {page_num}",
    #             content=page_text.strip(),
    #             source=source,
    #             metadata={**metadata, "page_number": page_num}
    #         )
    #         documents.append(page_doc)
    
    return documents


async def _parse_docx_file(file_path: Path, source: str) -> List[DocumentInput]:
    """
    DOCX 파일 파싱
    
    Args:
        file_path: 파일 경로
        source: 문서 출처
        
    Returns:
        List[DocumentInput]: 파싱된 문서
    """
    documents = []
    
    # DOCX 읽기
    doc = docx.Document(str(file_path))
    
    # 텍스트 추출
    full_text = ""
    paragraphs = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
            full_text += para.text + "\n\n"
    
    # 표에서 텍스트 추출
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                full_text += row_text + "\n"
        full_text += "\n"
    
    # 문서 ID 생성
    doc_id = _generate_document_id(file_path.name, full_text)
    
    # 메타데이터
    metadata = {
        "file_name": file_path.name,
        "file_size": file_path.stat().st_size,
        "paragraph_count": len(paragraphs),
        "table_count": len(doc.tables),
        "created_at": datetime.utcnow().isoformat()
    }
    
    # 문서 속성 추가
    if doc.core_properties.title:
        metadata["docx_title"] = doc.core_properties.title
    if doc.core_properties.author:
        metadata["docx_author"] = doc.core_properties.author
    if doc.core_properties.subject:
        metadata["docx_subject"] = doc.core_properties.subject
    
    # DocumentInput 생성
    document = DocumentInput(
        document_id=doc_id,
        title=metadata.get("docx_title", file_path.stem),
        content=full_text.strip(),
        source=source,
        metadata=metadata
    )
    
    documents.append(document)
    return documents


async def _parse_csv_file(file_path: Path, source: str) -> List[DocumentInput]:
    """
    CSV 파일 파싱
    
    Args:
        file_path: 파일 경로
        source: 문서 출처
        
    Returns:
        List[DocumentInput]: 파싱된 문서
    """
    documents = []
    
    # CSV 읽기
    df = pd.read_csv(file_path, encoding='utf-8')
    
    # 각 행을 별도 문서로 변환
    for idx, row in df.iterrows():
        # 행 데이터를 텍스트로 변환
        content_parts = []
        metadata = {
            "file_name": file_path.name,
            "row_index": idx,
            "source_type": "csv"
        }
        
        for col in df.columns:
            value = row[col]
            if pd.notna(value):
                content_parts.append(f"{col}: {value}")
                # 메타데이터에도 추가
                metadata[f"csv_{col}"] = str(value)
        
        content = "\n".join(content_parts)
        
        # 문서 ID 생성
        doc_id = _generate_document_id(f"{file_path.name}_row_{idx}", content)
        
        # DocumentInput 생성
        document = DocumentInput(
            document_id=doc_id,
            title=f"{file_path.stem} - Row {idx + 1}",
            content=content,
            source=source,
            metadata=metadata
        )
        
        documents.append(document)
    
    # 전체 요약 문서 추가
    summary_content = f"CSV 파일: {file_path.name}\n"
    summary_content += f"총 행 수: {len(df)}\n"
    summary_content += f"컬럼: {', '.join(df.columns)}\n"
    
    summary_doc = DocumentInput(
        document_id=_generate_document_id(file_path.name, "summary"),
        title=f"{file_path.stem} - Summary",
        content=summary_content,
        source=source,
        metadata={
            "file_name": file_path.name,
            "file_size": file_path.stat().st_size,
            "row_count": len(df),
            "column_count": len(df.columns),
            "columns": list(df.columns),
            "document_type": "csv_summary"
        }
    )
    
    documents.insert(0, summary_doc)
    return documents


async def _parse_json_file(file_path: Path, source: str) -> List[DocumentInput]:
    """
    JSON 파일 파싱
    
    Args:
        file_path: 파일 경로
        source: 문서 출처
        
    Returns:
        List[DocumentInput]: 파싱된 문서
    """
    documents = []
    
    # JSON 읽기
    async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
        content = await f.read()
        data = json.loads(content)
    
    # JSON 구조에 따라 다르게 처리
    if isinstance(data, list):
        # 배열인 경우 각 항목을 별도 문서로
        for idx, item in enumerate(data):
            item_content = json.dumps(item, ensure_ascii=False, indent=2)
            
            doc_id = _generate_document_id(f"{file_path.name}_item_{idx}", item_content)
            
            document = DocumentInput(
                document_id=doc_id,
                title=f"{file_path.stem} - Item {idx + 1}",
                content=item_content,
                source=source,
                metadata={
                    "file_name": file_path.name,
                    "item_index": idx,
                    "source_type": "json_array",
                    "item_type": type(item).__name__
                }
            )
            
            documents.append(document)
    
    else:
        # 단일 객체인 경우
        content_text = json.dumps(data, ensure_ascii=False, indent=2)
        
        doc_id = _generate_document_id(file_path.name, content_text)
        
        document = DocumentInput(
            document_id=doc_id,
            title=file_path.stem,
            content=content_text,
            source=source,
            metadata={
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
                "source_type": "json_object",
                "root_keys": list(data.keys()) if isinstance(data, dict) else []
            }
        )
        
        documents.append(document)
    
    return documents


def _generate_document_id(identifier: str, content: str) -> str:
    """
    문서 ID 생성
    
    Args:
        identifier: 식별자 (파일명 등)
        content: 문서 내용
        
    Returns:
        str: 생성된 문서 ID
    """
    # 해시 생성
    hash_input = f"{identifier}:{content[:1000]}"  # 처음 1000자만 사용
    hash_object = hashlib.sha256(hash_input.encode())
    hash_hex = hash_object.hexdigest()
    
    # 타임스탬프 추가
    timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
    
    return f"doc_{timestamp}_{hash_hex[:12]}"


async def detect_file_type(file_path: Path) -> Dict[str, Any]:
    """
    파일 타입 감지
    
    Args:
        file_path: 파일 경로
        
    Returns:
        Dict[str, Any]: 파일 타입 정보
    """
    # MIME 타입 감지
    mime_type, _ = mimetypes.guess_type(str(file_path))
    
    # 확장자
    extension = file_path.suffix.lower().lstrip('.')
    
    # 파일 크기
    file_size = file_path.stat().st_size
    
    # 지원 여부 확인
    supported_extensions = ['txt', 'pdf', 'docx', 'csv', 'json']
    is_supported = extension in supported_extensions
    
    return {
        "mime_type": mime_type,
        "extension": extension,
        "file_size": file_size,
        "file_size_mb": round(file_size / (1024 * 1024), 2),
        "is_supported": is_supported,
        "supported_extensions": supported_extensions
    }